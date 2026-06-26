from pathlib import Path
import re
import unicodedata
from zipfile import BadZipFile

from sqlalchemy.orm import Session

from app.core.config import UPLOAD_DIR
from app.models.project_document import ProjectDocument
from app.models.skill import Skill
from app.models.skill_alias import SkillAlias
from app.models.task import Task
from app.services.skills_service import list_skills

TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".xml", ".html", ".css", ".js", ".ts", ".tsx", ".py", ".java", ".sql"}
DOCUMENT_UPLOAD_DIR = Path(UPLOAD_DIR) / "documents"
MIN_CONFIDENCE = 0.7
MAX_DOCUMENT_CHARS = 20_000

STOPWORDS = {
    "task",
    "tasks",
    "project",
    "proiect",
    "feature",
    "functionalitate",
    "functie",
    "implementare",
    "implementa",
    "dezvoltare",
    "sistem",
    "aplicatie",
    "pagina",
    "modul",
}

BUILTIN_ALIASES = {
    "javascript": ["js", "ecmascript"],
    "typescript": ["ts"],
    "react": ["react.js", "reactjs"],
    "node.js": ["node", "nodejs"],
    "postgresql": ["postgres", "psql"],
    "mysql": ["mysql database"],
    "machine learning": ["ml", "model training"],
    "artificial intelligence": ["ai"],
    "user interface": ["ui", "interfata utilizator"],
    "user experience": ["ux"],
    "rest api": ["rest", "api rest"],
    "fastapi": ["fast api"],
}


def normalize_text(value: str) -> str:
    value = unicodedata.normalize("NFKD", value)
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    value = re.sub(r"[^a-z0-9+#.]+", " ", value.lower())
    value = re.sub(r"(?<=\w)\.(?=\s|$)", " ", value)
    return value.strip()


def phrase_pattern(phrase: str) -> str:
    normalized = normalize_text(phrase)
    escaped = re.escape(normalized).replace(r"\ ", r"\s+")
    return rf"(?<![a-z0-9+#.]){escaped}(?![a-z0-9+#.])"


def skill_terms(db: Session, skill: Skill) -> list[tuple[str, str, float]]:
    terms: list[tuple[str, str, float]] = [(skill.name, "nume skill", 0.95)]
    normalized_name = normalize_text(skill.name)

    for alias in BUILTIN_ALIASES.get(normalized_name, []):
        terms.append((alias, "alias tehnic", 0.92))

    aliases = db.query(SkillAlias).filter(SkillAlias.skill_id == skill.id).all()
    for alias in aliases:
        terms.append((alias.alias, "alias skillbook", 0.92))

    seen = set()
    unique_terms: list[tuple[str, str, float]] = []
    for term, source, confidence in terms:
        normalized = normalize_text(term)
        if not normalized or normalized in STOPWORDS or normalized in seen:
            continue
        seen.add(normalized)
        unique_terms.append((term, source, confidence))

    return unique_terms


def stored_document_path(document_id: int) -> Path | None:
    if not DOCUMENT_UPLOAD_DIR.exists():
        return None
    matches = sorted(DOCUMENT_UPLOAD_DIR.glob(f"{document_id}_*"))
    return matches[0] if matches else None


def read_document_text(document: ProjectDocument) -> str:
    pieces = [document.file_name, document.description or ""]
    path = stored_document_path(document.id)
    if path and path.exists():
        extracted = extract_file_text(path)
        if extracted:
            pieces.append(extracted[:MAX_DOCUMENT_CHARS])
    return "\n".join(piece for piece in pieces if piece)


def extract_file_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in TEXT_EXTENSIONS:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".pdf":
            return extract_pdf_text(path)
        if suffix == ".docx":
            return extract_docx_text(path)
    except OSError:
        return ""
    return ""


def extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""

    try:
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return ""


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""

    try:
        document = Document(str(path))
    except (BadZipFile, OSError, ValueError):
        return ""

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join([*paragraphs, *table_cells])


def build_task_corpus(db: Session, task: Task) -> tuple[str, int]:
    documents = (
        db.query(ProjectDocument)
        .filter(ProjectDocument.project_id == task.project_id, ProjectDocument.task_id == task.id)
        .all()
    )
    parts = [task.title, task.description or ""]
    parts.extend(read_document_text(document) for document in documents)
    return "\n".join(part for part in parts if part), len(documents)


def match_skill(db: Session, skill: Skill, normalized_corpus: str) -> tuple[bool, float, str, str | None]:
    best: tuple[float, str, str | None] = (0, "", None)

    for term, source, confidence in skill_terms(db, skill):
        normalized_term = normalize_text(term)
        if re.search(phrase_pattern(normalized_term), normalized_corpus):
            reason = f"Potrivire prin {source}: '{term}'"
            if confidence > best[0]:
                best = (confidence, reason, term)

    if best[0] >= MIN_CONFIDENCE:
        return True, best[0], best[1], best[2]

    return False, 0, "", None


def extract_task_skills(db: Session, task: Task) -> dict:
    corpus, document_count = build_task_corpus(db, task)
    normalized_corpus = normalize_text(corpus)
    suggestions = []

    for skill in list_skills(db):
        matched, confidence, reason, matched_term = match_skill(db, skill, normalized_corpus)
        if matched:
            suggestions.append(
                {
                    "skill_id": skill.id,
                    "name": skill.name,
                    "confidence": confidence,
                    "reason": reason,
                    "matched_term": matched_term,
                }
            )

    suggestions.sort(key=lambda item: (-item["confidence"], item["name"].lower()))
    return {
        "task_id": task.id,
        "document_count": document_count,
        "suggestions": suggestions,
    }
